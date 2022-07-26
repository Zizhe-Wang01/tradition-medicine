import os
import docx
#from win32com import client as wc
from docx import Document
from docx.shared import RGBColor, Pt
from openpyxl import load_workbook
import glob
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

'''
# 转化RTF文件为.docx文件存入temp_news目录
def transform_document(path, dest_path):
  word = wc.Dispatch('word.application')
  for file in os.listdir(path):
    (file_path, temp_file_name) = os.path.split(file)
    (short_name, extension) = os.path.splitext(temp_file_name)
    print(short_name)
    docx = word.Documents.Open(path + file)
    docx.SaveAs(dest_path + short_name + ".docx", 16)
    docx.Close()
  word.Quit()
'''




# 读取文件夹中所有的文件, 并返回存储所有文件路径的列表
def read_total_files(path):
  g = os.walk(path)
  files_list = []
  for path, dir_list, file_list in g:
    for file_name in file_list:
        files_list.append(os.path.join(path, file_name))
  print("read file list successfully")
  return files_list



# 整理title的
def get_title(files_list):
    index = ''
    news_dir_title = {}
    code_list = []
    for each in files_list:
      doc = docx.Document(each)
      for p in doc.paragraphs[0:]:
         for r in p.runs:
             if r.font.color.rgb == RGBColor(255, 0, 0) and r.font.color.rgb != \
                     RGBColor(91, 155, 213):
                 index = r.text
                 code_list.append(index)
                 if index not in news_dir_title.keys():
                     news_dir_title[index] = ""
             if r.font.size == Pt(14) :
                    temp = ''
                    temp += r.text
                    news_dir_title[index] += temp
    print("read title successfully")
    #print(code_list)
    return code_list, news_dir_title


# 整理content的
def get_content(files_list):
  count = ''
  news_dir_content = {}
  title_is = {}
  tag = True
  for each in files_list:
      doc = docx.Document(each)
      for p in doc.paragraphs[0:]:
         tag = True
         if p.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
             tag = True
             for r in p.runs:
                 if r.font.color.rgb == RGBColor(255, 0, 0):
                     tag = False
                     count = r.text
                     if count not in news_dir_content.keys():
                         news_dir_content[count] = ""
                         title_is[count] = False
                 elif r.font.size == Pt(14) and count in title_is.keys() and r.font.bold == True:
                      title_is[count] = True
                      tag = False
                      break
                 elif p.runs[0].font.bold == True and p.runs[0].font.color.rgb != RGBColor(204, 0, 51):
                     tag = False
                 elif (p.runs[0].font.size == Pt(10)) or p.runs[0].font.color.rgb == RGBColor(204, 0, 51):
                     tag = True
             if tag == True and title_is[count] == True:
                 temp = ''
                 temp += p.text
                 news_dir_content[count] += temp
  print("get content successfully")
  return news_dir_content


#插入文章标号
def insert_article_code(out_path, in_path,code_list, news_dir_title, news_dir_content):
  workbook = load_workbook(out_path + in_path)
  sheet = workbook.active
  for i in range(2, len(code_list)+1):
    sheet['A'+str(i)] = code_list[i-2]
  #workbook.save(path + r'\Coding spreadsheet - Tone+frame.xlsx')

# #插入文章标题


  sheet['AA1'] = 'title'
  for i in range(2, len(news_dir_title.values())+1):
      content = news_dir_title[code_list[i-2]] + news_dir_content[code_list[i-2]]
      sheet['AA'+str(i)] = content
  workbook.save(out_path +r'\new_sheet.xlsx')

#插入文章内容
'''
  sheet['AB1'] = 'content'
  for i in range(2, len(news_dir_content.values())+1):
    sheet['AB'+str(i)] = news_dir_content[code_list[i-2]]
  workbook.save(out_path +r'\new_sheet.xlsx')   #保存excel文件
'''